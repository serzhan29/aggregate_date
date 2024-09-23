from django.contrib import admin
from import_export import resources, fields
from import_export.admin import ExportActionModelAdmin
from .models import User, Indicator, TeacherReport, AdminReport
from import_export.formats.base_formats import XLSX
from django.http import HttpResponse
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from import_export.widgets import ForeignKeyWidget


# Настройка админки для кастомной модели User
@admin.register(User)
class UserAdmin(admin.ModelAdmin):
    list_display = ('username', 'email', 'role')  # Поля, которые будут отображаться в списке
    search_fields = ('username', 'email')  # Поля, по которым можно искать
    list_filter = ('role',)  # Фильтрация по полям

# Настройка админки для модели Indicator
@admin.register(Indicator)
class IndicatorAdmin(admin.ModelAdmin):
    list_display = ('name', 'unit')  # Поля, которые будут отображаться в списке
    search_fields = ('name',)  # Поля, по которым можно искать

# Настройка админки для модели TeacherReport
@admin.register(TeacherReport)
class TeacherReportAdmin(admin.ModelAdmin):
    list_display = ('teacher', 'indicator', 'plan_2022_2023', 'actual_2023_2024', 'plan_2024_2025', 'comment')
    search_fields = ('teacher__username', 'indicator__name')  # Поля, по которым можно искать
    list_filter = ('teacher', 'indicator')  # Фильтрация по полям

# Ресурс для модели AdminReport
class AdminReportResource(resources.ModelResource):
    indicator = fields.Field(
        column_name='Индикатор',
        attribute='indicator',
        widget=ForeignKeyWidget(Indicator, 'name')
    )
    total_plan_2022_2023 = fields.Field(
        column_name='Суммарный план 2022-2023',
        attribute='total_plan_2022_2023'
    )
    total_actual_2023_2024 = fields.Field(
        column_name='Суммарные выполненные показатели 2023-2024',
        attribute='total_actual_2023_2024'
    )
    total_plan_2024_2025 = fields.Field(
        column_name='Суммарный план на 2024-2025',
        attribute='total_plan_2024_2025'
    )

    class Meta:
        model = AdminReport
        fields = ('indicator', 'total_plan_2022_2023', 'total_actual_2023_2024', 'total_plan_2024_2025')

# Определение админки для модели AdminReport
@admin.register(AdminReport)
class AdminReportAdmin(ExportActionModelAdmin):
    resource_class = AdminReportResource
    list_display = ('indicator_name', 'total_plan_2022_2023', 'total_actual_2023_2024', 'total_plan_2024_2025')
    search_fields = ('indicator__name',)
    formats = [XLSX]

    def indicator_name(self, obj):
        return obj.indicator.name
    indicator_name.short_description = 'Индикатор'

    def get_actions(self, request):
        actions = super().get_actions(request)
        actions['export_to_pdf'] = (self.export_to_pdf, 'export_to_pdf', 'Export selected reports to PDF')
        actions['export_to_word'] = (self.export_to_word, 'export_to_word', 'Export selected reports to Word')
        return actions

    def export_to_pdf(self, request, queryset):
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        height -= 50

        for obj in queryset:
            p.drawString(100, height, f"Indicator: {obj.indicator.name}")
            p.drawString(100, height - 20, f"Total Plan 2022-2023: {obj.total_plan_2022_2023}")
            p.drawString(100, height - 40, f"Total Actual 2023-2024: {obj.total_actual_2023_2024}")
            p.drawString(100, height - 60, f"Total Plan 2024-2025: {obj.total_plan_2024_2025}")
            height -= 100

        p.showPage()
        p.save()

        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="admin_report.pdf"'
        return response

    def export_to_word(self, request, queryset):
        doc = Document()
        doc.add_heading('Admin Report', 0)

        for obj in queryset:
            doc.add_paragraph(f"Indicator: {obj.indicator.name}")
            doc.add_paragraph(f"Total Plan 2022-2023: {obj.total_plan_2022_2023}")
            doc.add_paragraph(f"Total Actual 2023-2024: {obj.total_actual_2023_2024}")
            doc.add_paragraph(f"Total Plan 2024-2025: {obj.total_plan_2024_2025}")
            doc.add_paragraph()

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=admin_report.docx'
        doc.save(response)
        return response
